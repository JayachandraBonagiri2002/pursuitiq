"""
export_proposal.py — Generate downloadable DOCX proposal from pursuit results.

Produces a professional Word document that bid teams can immediately edit,
brand, and submit to clients. Includes all sections, pricing summary,
and architecture diagram placeholder.
"""

import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def generate_proposal_docx(pursuit: dict) -> bytes:
    """
    Generate a professional DOCX proposal from pursuit data.

    Returns bytes of the .docx file ready for download.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    decomp = pursuit.get("decomposition", {})
    draft = pursuit.get("draft", {})
    pricing_data = pursuit.get("solution_pricing", {})
    comp = pursuit.get("competitor", {})
    win = pursuit.get("win_intel", {})
    client = pursuit.get("client_intel", {})

    # ── Cover Page ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROPOSAL")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(88, 28, 135)
    run.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(decomp.get("title", "Proposal"))
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()

    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = client_para.add_run(f"Prepared for: {decomp.get('client_name', 'Client')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details = []
    if decomp.get("industry"):
        details.append(f"Industry: {decomp['industry']}")
    if decomp.get("geography"):
        details.append(f"Geography: {', '.join(decomp['geography'])}")
    if decomp.get("estimated_deal_size_usd"):
        details.append(f"Deal Size: {decomp['estimated_deal_size_usd']}")
    run = meta.add_run(" | ".join(details))
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_page_break()

    # ── Table of Contents placeholder ─────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    toc_items = ["Executive Summary", "Understanding Your Challenge",
                 "Our Proposed Solution", "Delivery Approach and Methodology",
                 "Team and Credentials", "Why Choose Us",
                 "Pricing Summary", "Appendix: Competitive Positioning"]
    for i, item in enumerate(toc_items, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")
    doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    if draft.get("executive_summary"):
        doc.add_paragraph(draft["executive_summary"])
    doc.add_paragraph()

    # Win themes
    if draft.get("win_themes"):
        doc.add_heading("Our Key Differentiators", level=2)
        for theme in draft["win_themes"]:
            doc.add_paragraph(theme, style="List Bullet")
    doc.add_page_break()

    # ── Proposal Sections ─────────────────────────────────────────────────────
    if draft.get("sections"):
        for section in draft["sections"]:
            doc.add_heading(section.get("section_title", "Section"), level=1)
            content = section.get("content", "")
            for para in content.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_page_break()

    # ── Pricing Summary ───────────────────────────────────────────────────────
    doc.add_heading("Pricing Summary", level=1)

    if pricing_data.get("pricing"):
        p = pricing_data["pricing"]
        doc.add_heading("Investment Overview", level=2)

        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Shading Accent 1"
        cells = table.rows[0].cells
        cells[0].text = "Recommended Price"
        cells[1].text = f"${p.get('recommended_price_usd', 0):,.0f}"
        cells = table.rows[1].cells
        cells[0].text = "Pricing Structure"
        cells[1].text = p.get("pricing_structure", "")
        cells = table.rows[2].cells
        cells[0].text = "Price Range (Low)"
        cells[1].text = f"${p.get('price_low_usd', 0):,.0f}"
        cells = table.rows[3].cells
        cells[0].text = "Price Range (High)"
        cells[1].text = f"${p.get('price_high_usd', 0):,.0f}"
        cells = table.rows[4].cells
        cells[0].text = "Confidence Level"
        cells[1].text = f"{p.get('confidence', 0) * 100:.0f}%"

        doc.add_paragraph()

    # Solution options
    if pricing_data.get("solution_options"):
        doc.add_heading("Solution Options", level=2)
        for opt in pricing_data["solution_options"]:
            is_rec = opt.get("recommended", False)
            prefix = "[RECOMMENDED] " if is_rec else ""
            doc.add_heading(f"{prefix}{opt.get('name', 'Option')}", level=3)
            doc.add_paragraph(opt.get("description", ""))
            doc.add_paragraph(f"Total Investment: ${opt.get('total_cost_usd', 0):,.0f}")
            doc.add_paragraph(f"Delivery Timeline: {opt.get('delivery_months', 0)} months")
            doc.add_paragraph(f"Risk Level: {opt.get('risk_level', 'medium')}")
            if opt.get("key_components"):
                doc.add_paragraph("Key Components:")
                for comp_item in opt["key_components"]:
                    doc.add_paragraph(comp_item, style="List Bullet")
            doc.add_paragraph()

    doc.add_page_break()

    # ── Appendix: Architecture Diagram ────────────────────────────────────────
    if draft.get("architecture_diagram"):
        doc.add_heading("Appendix: Solution Architecture", level=1)
        doc.add_paragraph(
            "The following Mermaid diagram represents the proposed solution architecture. "
            "Render at mermaid.live or include as an image in the final submission."
        )
        doc.add_paragraph()
        arch_para = doc.add_paragraph()
        run = arch_para.add_run(draft["architecture_diagram"])
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    # ── Appendix: Competitive Positioning ─────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Appendix: Competitive Positioning (Internal Only)", level=1)
    doc.add_paragraph(
        "NOTE: This section is for internal use only. Remove before client submission.",
    )
    doc.add_paragraph()

    if comp.get("killer_differentiator"):
        doc.add_heading("Killer Differentiator", level=2)
        doc.add_paragraph(comp["killer_differentiator"])

    if comp.get("competitors"):
        doc.add_heading("Competitor Analysis", level=2)
        for c in comp["competitors"]:
            doc.add_heading(c.get("competitor_name", ""), level=3)
            doc.add_paragraph(f"Likelihood to bid: {c.get('likelihood_to_bid', '')}")
            doc.add_paragraph(f"Predicted positioning: {c.get('predicted_positioning', '')}")
            doc.add_paragraph(f"Price range: {c.get('predicted_price_range_usd', '')}")
            if c.get("how_to_beat_them"):
                doc.add_paragraph("How to beat them:")
                for tactic in c["how_to_beat_them"]:
                    doc.add_paragraph(tactic, style="List Bullet")
            doc.add_paragraph()

    # ── Save to bytes ─────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
